import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from openpyxl import load_workbook
import os
import re
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Function definitions
def apply_font_formatting(word_cell, excel_font):
    try:
        run = word_cell.paragraphs[0].runs[0]
        if excel_font.name:
            run.font.name = excel_font.name
        if excel_font.size:
            run.font.size = Pt(excel_font.size)
        run.font.bold = excel_font.bold
        run.font.italic = excel_font.italic
        underline_mapping = {
            'single': WD_UNDERLINE.SINGLE,
            'double': WD_UNDERLINE.DOUBLE,
            'none': None
        }
        underline_value = underline_mapping.get(excel_font.underline, None)
        run.font.underline = underline_value
    except Exception as e:
        print(f"Error applying font formatting: {e}")

def myreplace(doc, t, replace):
    try:
        replace = str(replace)
        for p in doc.paragraphs:
            if t.search(p.text):
                inline = p.runs
                for i in range(len(inline)):
                    if t.search(inline[i].text):
                        text = t.sub(replace, inline[i].text)
                        inline[i].text = text
    except Exception as e:
        print(f"Error in myreplace function: {e}")

def apply_borders(word_cell, excel_cell):
    borders = excel_cell.border
    if borders:
        # Apply top border
        if borders.top.style:
            word_cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s><w:top w:val="single" w:sz="4"/></w:tcBorders>' % nsdecls('w')))
        # Apply bottom border
        if borders.bottom.style:
            word_cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s><w:bottom w:val="single" w:sz="4"/></w:tcBorders>' % nsdecls('w')))
        # Apply left border
        if borders.left.style:
            word_cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s><w:left w:val="single" w:sz="4"/></w:tcBorders>' % nsdecls('w')))
        # Apply right border
        if borders.right.style:
            word_cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s><w:right w:val="single" w:sz="4"/></w:tcBorders>' % nsdecls('w')))

def insert_table_from_range(doc, table_sheet, start_cell, end_cell, tag):
    try:
        start_col, start_row = re.match(r"([A-Z]+)(\d+)", start_cell).groups()
        end_col, end_row = re.match(r"([A-Z]+)(\d+)", end_cell).groups()
        start_col_idx = ord(start_col) - ord('A')
        end_col_idx = ord(end_col) - ord('A')
        start_row_idx = int(start_row) - 1
        end_row_idx = int(end_row) - 1
        num_rows = end_row_idx - start_row_idx + 1
        num_cols = end_col_idx - start_col_idx + 1
        for p in doc.paragraphs:
            if re.search(re.escape(tag), p.text):
                p.clear()
                table = doc.add_table(rows=num_rows, cols=num_cols)
                for i in range(num_rows):
                    for j in range(num_cols):
                        excel_cell = table_sheet.cell(row=start_row_idx + i + 1, column=start_col_idx + j + 1)
                        word_cell = table.cell(i, j)
                        cell_value = str(excel_cell.value).replace('None', '') if excel_cell.value is not None else ''
                        word_cell.text = cell_value
                        word_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        apply_font_formatting(word_cell, excel_cell.font)
                        apply_borders(word_cell, excel_cell)
                p._element.addnext(table._element)
                break
    except Exception as e:
        print(f"Error inserting table from range: {e}")

def handle_replacements(doc, sheet1, workbook):
    try:
        for row in sheet1.iter_rows(min_row=2, max_row=sheet1.max_row, min_col=1, max_col=4, values_only=True):
            tag, tag_type, replacement, table_range = row
            if tag is None or tag_type is None:
                continue
            if tag_type.lower() == 'word':
                t = re.compile(re.escape(tag))
                myreplace(doc, t, replacement)
            elif tag_type.lower() == 'table' and table_range is not None:
                if table_range.startswith('='):
                    table_range = table_range[1:]
                match = re.match(r"(\w+)!([A-Z]+\d+):([A-Z]+\d+)", table_range)
                if match:
                    sheet_name, start_cell, end_cell = match.groups()
                    print(f"Table range detected: {sheet_name} {start_cell}:{end_cell}")
                    table_sheet = workbook[sheet_name]
                    insert_table_from_range(doc, table_sheet, start_cell, end_cell, tag)
    except Exception as e:
        print(f"Error handling replacements: {e}")

def browse_word_file():
    filename = filedialog.askopenfilename(title="Select Word File", filetypes=(("Word files", "*.docx"), ("All files", "*.*")))
    word_file_path.set(filename)

def browse_excel_file():
    filename = filedialog.askopenfilename(title="Select Excel File", filetypes=(("Excel files", "*.xlsx"), ("All files", "*.*")))
    excel_file_path.set(filename)

def run_automation():
    try:
        doc_filename = word_file_path.get()
        excel_filename = excel_file_path.get()

        if not os.path.exists(doc_filename):
            messagebox.showerror("Error", f"Document file {doc_filename} does not exist.")
            return
        if not os.path.exists(excel_filename):
            messagebox.showerror("Error", f"Excel file {excel_filename} does not exist.")
            return

        doc = Document(doc_filename)
        workbook = load_workbook(excel_filename)
        sheet1 = workbook['Sheet1']

        handle_replacements(doc, sheet1, workbook)

        doc.save(doc_filename)
        messagebox.showinfo("Success", "Document automation completed successfully.")
    except Exception as e:
        messagebox.showerror("Error", f"Error in main execution: {e}")

# Tkinter GUI
root = tk.Tk()
root.title("Document Automation")
root.attributes('-fullscreen', True)
root.configure(bg='#86bd24')  # Deloitte green

heading = tk.Label(root, text="Let's Automate", font=("Helvetica", 36), fg="white", bg="#86bd24")
heading.pack(pady=50)

word_file_path = tk.StringVar()
excel_file_path = tk.StringVar()

word_file_frame = tk.Frame(root, bg="#86bd24")
word_file_frame.pack(pady=20)
word_file_label = tk.Label(word_file_frame, text="Select Word File:", font=("Helvetica", 18), fg="white", bg="#86bd24")
word_file_label.pack(side="left", padx=10)
word_file_entry = tk.Entry(word_file_frame, textvariable=word_file_path, font=("Helvetica", 18), width=50)
word_file_entry.pack(side="left", padx=10)
word_file_button = tk.Button(word_file_frame, text="Browse", font=("Helvetica", 18), command=browse_word_file)
word_file_button.pack(side="left", padx=10)

excel_file_frame = tk.Frame(root, bg="#86bd24")
excel_file_frame.pack(pady=20)
excel_file_label = tk.Label(excel_file_frame, text="Select Excel File:", font=("Helvetica", 18), fg="white", bg="#86bd24")
excel_file_label.pack(side="left", padx=10)
excel_file_entry = tk.Entry(excel_file_frame, textvariable=excel_file_path, font=("Helvetica", 18), width=50)
excel_file_entry.pack(side="left", padx=10)
excel_file_button = tk.Button(excel_file_frame, text="Browse", font=("Helvetica", 18), command=browse_excel_file)
excel_file_button.pack(side="left", padx=10)

run_button = tk.Button(root, text="Run Automation", font=("Helvetica", 20), command=run_automation)
run_button.pack(pady=20)

footer_frame = tk.Frame(root, bg="#86bd24")
footer_frame.pack(side="bottom", fill="x", pady=10)

info_label = tk.Label(footer_frame, text="* In case you run into a permission denied problem, check if your Word file is open.", font=("Helvetica", 12), fg="grey", bg="#86bd24")
info_label.pack(side="left", padx=20)

root.mainloop()
